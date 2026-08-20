import io
import os

from django.conf import settings
from django.test import TestCase
from docx import Document

from review import models as review_models
from review.logic import serve_review_file
from utils.testing import helpers


class ServeReviewFileTests(TestCase):
    """
    Regression coverage for review.logic.serve_review_file, written ahead
    of the python-docx 0.8.11 -> 1.2.0 upgrade (Phase 1 of the Django 5.2
    migration plan) so we have a "before" baseline for the generated docx
    structure. Prior to this file, serve_review_file had zero test
    coverage.
    """

    @classmethod
    def setUpTestData(cls):
        cls.journal, _journal_two = helpers.create_journals()

        cls.review_form = helpers.create_review_form(cls.journal)
        cls.element_one = review_models.ReviewFormElement.objects.create(
            name="Summary",
            kind="text",
            order=1,
            required=True,
        )
        cls.element_two = review_models.ReviewFormElement.objects.create(
            name="Recommendation",
            kind="text",
            order=2,
            required=True,
        )
        cls.review_form.elements.add(cls.element_one, cls.element_two)

        cls.assignment = helpers.create_review_assignment(
            journal=cls.journal,
            review_form=cls.review_form,
        )
        cls.assignment.article.title = "A Sample Article for Review"
        cls.assignment.article.save()

    def test_serve_review_file_produces_valid_docx(self):
        # serve_review_file writes to BASE_DIR/files/temp, which is
        # gitignored and normally created at install/deploy time rather
        # than being present in a fresh checkout - ensure it exists so the
        # test is hermetic.
        os.makedirs(os.path.join(settings.BASE_DIR, "files", "temp"), exist_ok=True)

        response = serve_review_file(self.assignment)

        # The view returns a StreamingHttpResponse wrapping a FileWrapper
        # around the generated .docx; consume it fully to get the raw
        # document bytes back for inspection.
        docx_bytes = b"".join(response.streaming_content)
        self.assertGreater(len(docx_bytes), 0)

        document = Document(io.BytesIO(docx_bytes))
        heading_styles = ("Title", "Heading 1", "Heading 2", "Heading 3")
        headings = [
            paragraph.text
            for paragraph in document.paragraphs
            if paragraph.style.name in heading_styles
        ]

        # add_heading("Review #{pk}", 0) is the document title.
        self.assertEqual(headings[0], f"Review #{self.assignment.pk}")

        # A level-1 heading mentions the article title.
        self.assertTrue(
            any(self.assignment.article.title in heading for heading in headings)
        )

        # Every ReviewFormElement on the assignment's form gets its own
        # level-2 heading.
        for element in self.assignment.form.elements.all():
            self.assertIn(element.name, headings)
