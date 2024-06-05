__copyright__ = "Copyright 2017 Birkbeck, University of London"
__author__ = "Martin Paul Eve & Andy Byers"
__license__ = "AGPL v3"
__maintainer__ = "Birkbeck Centre for Technology and Publishing"

import csv
from datetime import timedelta
from uuid import uuid4
import os
import re

from django.conf import settings
from django.contrib import messages
from django.db.models import (
    Avg,
    Count,
    IntegerField,
    OuterRef,
    Prefetch,
    Subquery,
    Case,
    When,
    BooleanField,
    Value,
)
from django.shortcuts import redirect, reverse
from django.utils import timezone
from django.db import IntegrityError
from django.utils.safestring import mark_safe
from docx import Document

from utils import render_template, setting_handler, notify_helpers
from core import(
    files,
    email,
    models as core_models,
)
from review import models
from review.const import EditorialDecisions as ED
from events import logic as event_logic
from submission import models as submission_models


def get_reviewers(article, candidate_queryset, exclude_pks):
    prefetch_review_assignment = Prefetch(
        'reviewer',
        queryset=models.ReviewAssignment.objects.filter(
            article__journal=article.journal
        ).exclude(date_complete__isnull=True).order_by("-date_complete")
    )
    active_reviews_count = models.ReviewAssignment.objects.filter(
        is_complete=False,
        reviewer=OuterRef("id"),
    ).values(
        "reviewer_id",
    ).annotate(
        rev_count=Count("reviewer_id"),
    ).values("rev_count")

    rating_average = models.ReviewerRating.objects.filter(
        assignment__article__journal=article.journal,
        assignment__reviewer=OuterRef("id"),
    ).values(
        # Without this .values call, results are grouped by assignment...
        "assignment__article__journal",
    ).annotate(
        rating_average=Avg("rating"),
    ).values("rating_average")

    completed_reviewer_pks_subquery = article.completed_reviews_with_decision.values_list(
        'reviewer__pk',
        flat=True,
    )

    # TODO swap the below subqueries with filtered annotations on Django 2.0+
    reviewers = candidate_queryset.exclude(
        pk__in=exclude_pks,
    ).prefetch_related(
        prefetch_review_assignment,
        'interest',
    ).annotate(
        active_reviews_count=Subquery(
            active_reviews_count,
            output_field=IntegerField(),
        )
    ).annotate(
        rating_average=Subquery(rating_average, output_field=IntegerField()),
        is_past_reviewer=Case(
            When(pk__in=Subquery(completed_reviewer_pks_subquery),
                 then=True),
            default=False,
            output_field=BooleanField(),
        ),
    )

    if article.journal.get_setting('general', 'enable_suggested_reviewers'):
        article_keywords = [keyword.word for keyword in article.keywords.all()]
        reviewers = reviewers.annotate(
            is_suggested_reviewer=Case(
                When(interest__name__in=article_keywords,
                     then=Value(True)),
                default=Value(False),
                output_field=BooleanField(),
            )
        )

    return reviewers


def get_reviewer_candidates(article, user=None, reviewers_to_exclude=None):
    """ Builds a queryset of candidates for peer review for the given article
    :param article: an instance of submission.models.Article
    :param user: The user requesting candidates who would be filtered out
    :param reviewers_to_exclude: queryset of Account objects
    """
    reviewer_pks_to_exclude = [
        review.reviewer.pk for review in article.reviewassignment_set.filter(
            review_round=article.current_review_round_object(),
        )
    ]
    if user:
        reviewer_pks_to_exclude.append(user.pk)

    if reviewers_to_exclude:
        for reviewer in reviewers_to_exclude:
            reviewer_pks_to_exclude.append(
                reviewer.pk,
            )

    return get_reviewers(
        article,
        article.journal.users_with_role('reviewer'),
        reviewer_pks_to_exclude
    )


def get_suggested_reviewers(article, reviewers):
    suggested_reviewers = []
    keywords = [keyword.word for keyword in article.keywords.all()]
    for reviewer in reviewers:
        interests = [interest.name for interest in reviewer.interest.all()]
        for interest in interests:
            if interest in keywords:
                suggested_reviewers.append(reviewer)
                break

    return suggested_reviewers


def get_previous_round_reviewers(article):
    """
    Builds a queryset of candidates who have previously completed a review
    for this article.
    :param article: an instance of submission.models.Article
    """
    review_assignments = article.reviewassignment_set.filter(
        review_round=article.current_review_round_object())
    reviewer_pks_to_exclude = [
        review.reviewer.pk for review in review_assignments
    ]

    completed_review_reviewer_pks = [
        review.reviewer.pk for review in article.completed_reviews_with_decision
    ]
    return get_reviewers(
        article,
        core_models.Account.objects.filter(
            pk__in=completed_review_reviewer_pks,
        ),
        reviewer_pks_to_exclude,
    )


def get_assignment_context(request, article, editor, assignment):
    review_in_review_url = request.journal.site_url(
        reverse(
            'review_in_review',
            kwargs={'article_id': article.pk}
        )
    )
    email_context = {
        'article': article,
        'editor': editor,
        'assignment': assignment,
        'review_in_review_url': review_in_review_url,
    }

    return email_context



def get_review_url(request, review_assignment):
    review_url = request.journal.site_url(path=reverse(
            'do_review', kwargs={'assignment_id': review_assignment.id}
    ))

    access_codes = setting_handler.get_setting(
        'general',
        'enable_one_click_access',
        request.journal,
    ).value

    if access_codes:
        review_url = "{0}?access_code={1}".format(
            review_url,
            review_assignment.access_code,
        )

    return review_url


def get_article_details_for_review(article):
    detail_string = """
        <b>Article Details:</b><br />
        <b>Title</b>: {article.title}<br />
        <b>Section</b>: {section}<br />
        <b>Keywords</b>: {keywords}<br />
        <b>Abstract</b>:<br />
            {article.abstract}<br />
        """.format(
        article=article,
        section=article.section.name if article.section else None,
        keywords=", ".join(kw.word for kw in article.keywords.all()),
    )
    return mark_safe(detail_string)


def get_reviewer_notification_context(
    request, article, editor,
    review_assignment,
):
    review_url = get_review_url(request, review_assignment)
    article_details = get_article_details_for_review(article)

    email_context = {
        'article': article,
        'editor': editor,
        'review_assignment': review_assignment,
        'review_url': review_url,
        'article_details': article_details,
    }

    return email_context


def get_reviewer_notification(
    request, article, editor, review_assignment,
    reminder=False,
):

    email_context = get_reviewer_notification_context(
        request, article, editor, review_assignment,
    )
    if reminder and reminder == 'request':
        return render_template.get_message_content(
            request,
            email_context,
            'default_review_reminder'
        )
    elif reminder and reminder == 'accepted':
        return render_template.get_message_content(
            request,
            email_context,
            'accepted_review_reminder'
        )
    else:
        return render_template.get_message_content(
            request,
            email_context,
            'review_assignment'
        )


def get_withdrawal_notification_context(request, review_assignment):

    email_context = {
        'article': review_assignment.article,
        'review_assignment': review_assignment,
        'editor': request.user,
    }
    return email_context


def get_unassignment_context(request, assignment):
    email_context = {
        'article': assignment.article,
        'assignment': assignment,
        'editor': request.user,
    }

    return email_context


def get_decision_context(request, article, decision, author_review_url):
    email_context = {
        'article': article,
        'decision': decision,
        'review_url': author_review_url,
    }

    return email_context


def get_decision_content(request, article, decision, author_review_url):
    context = get_decision_context(
        request,
        article,
        decision,
        author_review_url,
    )
    template_name = "review_decision_{0}".format(decision)
    return render_template.get_message_content(request, context, template_name)


def get_revision_request_content(request, article, revision, draft=False):
    email_context = {
        'article': article,
        'revision': revision,
    }

    if not draft:
        do_revisions_url = request.journal.site_url(path=reverse(
            'do_revisions',
            kwargs={
                'article_id': article.pk,
                'revision_id': revision.pk,
            }
        ))
        email_context['do_revisions_url'] = do_revisions_url
    else:
        email_context['do_revisions_url'] = "{{ do_revisions_url }}"

    return render_template.get_message_content(request, email_context, 'request_revisions')


def get_share_review_content(request, article, review):
    url = request.journal.site_url(
        reverse(
            'reviewer_share_reviews',
            kwargs={'article_id': article.pk}
        )
    )
    email_context = {
        'article': article,
        'review': review,
        'url': url,
    }

    return render_template.get_message_content(
        request,
        email_context,
        'share_reviews_notification',
    )


def send_review_share_message(request, article, subject, form_data):
    for review_pk, email_content in form_data.items():
        review = models.ReviewAssignment.objects.get(
            pk=review_pk,
            article__journal=article.journal,
        )
        notify_helpers.send_email_with_body_from_user(
            request,
            subject,
            review.reviewer.email,
            email_content,
            log_dict={
                'level': 'Info',
                'action_text': f'Reviews link shared with {review.reviewer.full_name()}',
                'types': 'Review Sharing',
                'actor': request.user,
                'target': article,
            }
        )


def get_reviewer_from_post(request):
    reviewer_id = request.POST.get('reviewer')

    if reviewer_id:
        reviewer = core_models.Account.objects.get(pk=reviewer_id)

        # if this user is not a reviewer, return None to force an error on the form.
        if not reviewer.is_reviewer(request):
            return None

        return reviewer
    else:
        return None


def log_revision_event(text, user, revision_request):
    action = models.RevisionAction.objects.create(
        text=text,
        logged=timezone.now(),
        user=user
    )

    revision_request.actions.add(action)


def get_draft_email_message(request, article):
    review_in_review_url = request.journal.site_url(
        path=reverse(
            'review_draft_decision', args=[article.pk]
        )
    )
    email_context = {
        'article': article,
        'review_in_review_url': review_in_review_url,
    }

    return render_template.get_message_content(request, email_context, 'draft_message')


def group_files(article, reviews):
    files = list()

    for file in article.manuscript_files.all():
        files.append(file)

    for file in article.data_figure_files.all():
        files.append(file)

    for review in reviews:
        if review.for_author_consumption and review.display_review_file:
            files.append(review.review_file)

    return files


def handle_draft_declined(article, draft_decision, request):
    """
    Raises an event when a draft decision is declined by an editor.
    """
    kwargs = {
        'article': article,
        'request': request,
        'draft_decision': draft_decision,
        'skip': False,
    }

    draft_decision.editor_decline_rationale = request.POST.get(
        'editor_decline_rationale',
        '<p>Editor provided no rationale.</p>'
    )
    draft_decision.save()

    return event_logic.Events.raise_event(
        event_logic.Events.ON_DRAFT_DECISION_DECLINED,
        task_object=article,
        **kwargs,
    )


def handle_decision_action(article, draft, request):
    # Setup email data
    subject = 'Decision'
    subject_setting_name = None
    user_message = request.POST.get('email_message', 'No message found.')
    if draft.decision == ED.ACCEPT.value:
        subject_setting_name = 'subject_review_decision_accept'
    elif draft.decision == ED.DECLINE.value:
        subject_setting_name = 'subject_review_decision_decline'
    elif draft.decision == ED.UNDECLINE.value:
        subject_setting_name = 'subject_review_decision_undecline'
    if subject_setting_name:
        subject = setting_handler.get_email_subject_setting(
            setting_name=subject_setting_name,
            journal=request.journal,
        )
    email_data = email.EmailData(
        subject=subject,
        body=user_message,
    )

    kwargs = {
        'article': article,
        'request': request,
        'decision': draft.decision,
        'skip': False,
        'email_data': email_data,
    }

    if draft.decision == ED.ACCEPT.value:
        article.accept_article(stage=submission_models.STAGE_EDITOR_COPYEDITING)
        event_logic.Events.raise_event(
            event_logic.Events.ON_ARTICLE_ACCEPTED,
            task_object=article,
            **kwargs,
        )
        # Call workflow element complete event
        workflow_kwargs = {
            'handshake_url': 'review_home',
            'request': request,
            'article': article,
            'switch_stage': True,
        }
        return event_logic.Events.raise_event(
            event_logic.Events.ON_WORKFLOW_ELEMENT_COMPLETE,
            task_object=article,
            **workflow_kwargs,
        )
    elif draft.decision == ED.DECLINE.value:
        article.decline_article()
        event_logic.Events.raise_event(
            event_logic.Events.ON_ARTICLE_DECLINED,
            task_object=article,
            **kwargs,
        )
    elif draft.decision == 'minor_revisions' or draft.decision == 'major_revisions':
        revision = models.RevisionRequest.objects.create(
            article=article,
            editor=draft.section_editor,
            editor_note='',
            type=draft.decision,
            date_due=draft.revision_request_due_date
        )
        do_revisions_url = request.journal.site_url(path=reverse(
            'do_revisions',
            kwargs={
                'article_id': article.pk,
                'revision_id': revision.pk,
            }
        ))
        revision_rendered_template = render_template.get_message_content(
            request,
            {'do_revisions_url': do_revisions_url},
            user_message,
            template_is_setting=True,
        )
        article.stage = submission_models.STAGE_UNDER_REVISION
        article.save()

        kwargs['user_message_content'] = revision_rendered_template
        kwargs['revision'] = revision
        event_logic.Events.raise_event(
            event_logic.Events.ON_REVISIONS_REQUESTED_NOTIFY,
            **kwargs,
        )
        return redirect(
            reverse(
                'view_revision',
                kwargs={
                    'article_id': article.pk,
                    'revision_id': revision.pk,
                }
            )
        )

    return None


def get_access_code(request):
    if request.GET.get('access_code'):
        access_code = request.GET.get('access_code')
    else:
        access_code = None

    return access_code


def quick_assign(request, article, reviewer_user=None):
    errors = []
    try:
        default_review_form_id = setting_handler.get_setting('general',
                                                             'default_review_form',
                                                             request.journal).processed_value
    except models.ReviewForm.DoesNotExist:
        errors.append('This journal has no default review form.')

    try:
        review_form = models.ReviewForm.objects.get(pk=default_review_form_id)
    except ValueError:
        errors.append('Default review form is not an integer.')

    try:
        default_visibility = setting_handler.get_setting('general',
                                                         'default_review_visibility',
                                                         request.journal).value
        default_due = setting_handler.get_setting('general',
                                                  'default_review_days',
                                                  request.journal).value
    except Exception:
        errors.append('This journal does not have either default visibilty or default due.')

    if not reviewer_user:
        user_id = request.POST.get('quick_assign')
        user = core_models.Account.objects.get(pk=user_id)
    else:
        user = reviewer_user

    if user not in request.journal.users_with_role('reviewer'):
        errors.append('This user is not a reviewer for this journal.')

    if not errors:
        new_assignment = models.ReviewAssignment.objects.create(
            article=article,
            reviewer=user,
            editor=request.user,
            review_round=article.current_review_round_object(),
            form=review_form,
            access_code=uuid4(),
            visibility=default_visibility,
            date_due=timezone.now() + timedelta(days=int(default_due)),
        )

        article.stage = submission_models.STAGE_UNDER_REVIEW
        article.save()

        email_content = get_reviewer_notification(request, article, request.user, new_assignment)

        kwargs = {'user_message_content': email_content,
                  'review_assignment': new_assignment,
                  'request': request,
                  'skip': False,
                  'acknowledgement': False}

        event_logic.Events.raise_event(event_logic.Events.ON_REVIEWER_REQUESTED, **kwargs)
        event_logic.Events.raise_event(event_logic.Events.ON_REVIEWER_REQUESTED_ACKNOWLEDGE, **kwargs)

        return new_assignment

    else:
        for error in errors:
            messages.add_message(request, messages.WARNING, error)


def handle_reviewer_form(request, new_reviewer_form):
    account = new_reviewer_form.save(commit=False)
    account.is_active = True
    account.save()
    account.add_account_role('reviewer', request.journal)
    messages.add_message(request, messages.INFO, 'A new account has been created.')
    return account


def get_enrollable_users(request):
    account_roles = core_models.AccountRole.objects.filter(
        journal=request.journal,
        role__slug='reviewer',
    ).prefetch_related(
        'user',
    )
    users_with_role = [assignment.user.pk for assignment in account_roles]
    return core_models.Account.objects.all().order_by(
        'last_name',
    ).exclude(
        pk__in=users_with_role,
    )


def generate_access_code_url(url_name, assignment, access_code):

    reverse_url = reverse(url_name, kwargs={'assignment_id': assignment.pk})

    if access_code:
        reverse_url = '{reverse_url}?access_code={access_code}'.format(reverse_url=reverse_url, access_code=access_code)

    return reverse_url


def render_choices(choices):
    c_split = choices.split('|')
    return [(choice.capitalize(), choice) for choice in c_split]


def serve_review_file(assignment):
    """
    Produces a word document representing the review form.
    :param assignment: ReviewAssignment object
    :return: HttpStreamingResponse
    """
    elements = assignment.form.elements.all()
    document = Document()
    document.add_heading('Review #{pk}'.format(pk=assignment.pk), 0)
    document.add_heading('Review of `{article_title}`'.format(article_title=assignment.article.title),
                         level=1)
    document.add_paragraph()
    document.add_paragraph('Complete the form below, then upload it under the "FILE UPLOAD" section on your review page'
                           '. There is no need to complete the form on the web page if you are uploading this '
                           'document.')
    document.add_paragraph()

    for element in elements:
        document.add_heading(element.name, level=2)
        document.add_paragraph(element.help_text)
        if element.choices:
            choices = render_choices(element.choices)
            table = document.add_table(rows=1, cols=2)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Choice'
            hdr_cells[1].text = 'Indication'

            for choice in element.choices.split('|'):
                row_cells = table.add_row().cells
                row_cells[0].text = str(choice)
        document.add_paragraph()

    filename = '{uuid}.docx'.format(uuid=uuid4())
    filepath = os.path.join(settings.BASE_DIR, 'files', 'temp', filename)
    document.save(filepath)
    return files.serve_temp_file(filepath, filename)


def handle_review_file_switch(review, switch):
    if switch == 'true':
        review.display_review_file = True
    else:
        review.display_review_file = False

    review.save()


def get_reminder_content(reminder_type, article, review_assignment, request):
    """
    Fetches the right email for a reminder type
    """

    if reminder_type == 'request':
        return get_reviewer_notification(
            request,
            article,
            request.user,
            review_assignment,
            reminder=reminder_type,
        )
    elif reminder_type == 'accepted':
        return get_reviewer_notification(
            request,
            article,
            request.user,
            review_assignment,
            reminder=reminder_type,
        )


def send_review_reminder(request, form, review_assignment, reminder_type):
    """
    Sends a reminder email to a reviewer
    """
    desc = '{sender} sent review reminder of type {type} to {to}'.format(
        sender=request.user.full_name(),
        type=reminder_type,
        to=review_assignment.reviewer.full_name()
    )
    log_dict = {'level': 'Info',
                'action_text': desc,
                'types': 'Review Reminder',
                'target': review_assignment.article}

    notify_helpers.send_email_with_body_from_user(
        request,
        form.cleaned_data['subject'],
        review_assignment.reviewer.email,
        form.cleaned_data['body'],
        log_dict=log_dict
    )


def assign_editor(
        article,
        editor,
        assignment_type,
        request=None,
        skip=True,
        automate_email=False,
):
    from core.forms import SettingEmailForm
    assignment, created = models.EditorAssignment.objects.get_or_create(
        article=article,
        editor=editor,
        editor_type=assignment_type,
    )
    if request and created and automate_email:
        email_context = get_assignment_context(
            request,
            article,
            editor,
            assignment,
        )
        form = SettingEmailForm(
            setting_name="editor_assignment",
            email_context=email_context,
            request=request,
        )
        post_data = {
            'subject': form.fields['subject'].initial,
            'body': form.fields['body'].initial,
        }
        form = SettingEmailForm(
            post_data,
            setting_name="editor_assignment",
            email_context=email_context,
            request=request,
        )

        if form.is_valid():
            kwargs = {
                'email_data': form.as_dataclass(),
                'editor_assignment': assignment,
                'request': request,
                'skip': skip,
                'acknowledgement': False,
            }
            event_logic.Events.raise_event(
                event_logic.Events.ON_ARTICLE_ASSIGNED,
                task_object=article, **kwargs,
            )
            if not skip:
                event_logic.Events.raise_event(
                    event_logic.Events.ON_ARTICLE_ASSIGNED_ACKNOWLEDGE,
                    **kwargs,
                )
    return assignment, created


def process_reviewer_csv(path, request, article, form):
    """
    Iterates through a CSV c
    """
    try:
        csv_file = open(path, 'r', encoding="utf-8-sig")
        reader = csv.DictReader(csv_file)
        reviewers = []
        for row in reader:
            try:
                country = core_models.Country.objects.get(code=row.get('country'))
            except core_models.Country.DoesNotExist:
                country = None

            reviewer, created = core_models.Account.objects.get_or_create(
                email=row.get('email_address'),
                defaults={
                    'salutation': row.get('salutation', ''),
                    'first_name': row.get('firstname', ''),
                    'middle_name': row.get('middlename', ''),
                    'last_name': row.get('lastname', ''),
                    'department': row.get('department', ''),
                    'institution': row.get('institution', ''),
                    'country': country,
                    'is_active': True,
                }
            )

            try:
                review_interests = row.get('interests')
                re.split('[,;]+', review_interests)
            except (IndexError, AttributeError):
                review_interests = []

            for term in review_interests:
                interest, _ = core_models.Interest.objects.get_or_create(name=term)
                reviewer.interest.add(interest)

            # Add the reviewer role
            reviewer.add_account_role('reviewer', request.journal)

            review_assignment, c = models.ReviewAssignment.objects.get_or_create(
                article=article,
                reviewer=reviewer,
                review_round=article.current_review_round_object(),
                defaults={
                    'editor': request.user,
                    'date_due': form.cleaned_data.get('date_due'),
                    'form': form.cleaned_data.get('form'),
                    'visibility': form.cleaned_data.get('visibility'),
                    'access_code': uuid4(),
                }
            )
            review_url = get_review_url(request, review_assignment)
            html = render_template.get_message_content(
                request=request,
                context={
                    'article': article,
                    'editor': request.user,
                    'review_assignment': review_assignment,
                    'review_url': review_url,
                    'article_details': get_article_details_for_review(article),
                    'reason': row.get('reason')
                },
                template=form.cleaned_data.get('template'),
                template_is_setting=True,
            )

            # finally, call event
            kwargs = {'user_message_content': html,
                      'review_assignment': review_assignment,
                      'request': request,
                      'skip': False,
                      'acknowledgement': True}

            event_logic.Events.raise_event(
                event_logic.Events.ON_REVIEWER_REQUESTED_ACKNOWLEDGE,
                **kwargs,
            )

            reviewers.append(
                {
                    'account': reviewer,
                    'reason': row.get('reason'),
                    'review_assignment': review_assignment,
                }
            )
        return reviewers, None
    except (IntegrityError, IndexError) as e:
        return [], e


def handle_response_letter_upload(request, revision_request):
    response_letter = request.FILES.get('response_letter')
    if response_letter:
        file = files.save_file_to_article(
            response_letter,
            revision_request.article,
            owner=request.user,
            label='Response Letter',
        )
        revision_request.response_letter = file
        revision_request.save()
        return file

    messages.add_message(
        request,
        messages.WARNING,
        'No response letter selected.',
    )
    return


def get_distinct_reviews(reviews):
    """
    A utility function we have to use because MySQL does not support
    distinct('field'). Returns a list of review objects by distinct reviewer.
    """
    reviews_to_return = []
    reviewers = set()
    for review in reviews:
        if review.reviewer not in reviewers:
            reviews_to_return.append(review)
            reviewers.add(review.reviewer)

    return reviews_to_return
